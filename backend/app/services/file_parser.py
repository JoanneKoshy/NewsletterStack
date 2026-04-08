import base64
from pathlib import Path
from PyPDF2 import PdfReader
from docx import Document as DocxDocument


def extract_text_from_pdf(file_path: str) -> str:
    reader = PdfReader(file_path)
    text_parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            text_parts.append(text)
    return "\n\n".join(text_parts)


def extract_text_from_docx(file_path: str) -> str:
    doc = DocxDocument(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also grab text from tables (MIS reports often have tables)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                paragraphs.append(row_text)

    return "\n\n".join(paragraphs)


def pdf_to_base64(file_path: str) -> str:
    with open(file_path, "rb") as f:
        return base64.b64encode(f.read()).decode("utf-8")


def parse_file(file_path: str) -> str:
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    elif ext == ".txt":
        return Path(file_path).read_text()
    else:
        raise ValueError(f"Unsupported file type: {ext}")